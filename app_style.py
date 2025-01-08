import os
import requests
import streamlit as st
# from dotenv import load_dotenv
from anthropic import Anthropic
from docx import Document
from io import BytesIO
from openai import OpenAI
from docx.shared import Inches
import re
import json
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

st.set_page_config(page_title="Kids Screenplay Generator", page_icon="🎬", layout="wide")

# Load environment variables
# load_dotenv()
# ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  # Load OpenAI API key

# Initialize the Anthropic client
open_ai_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
anthropic_client = Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

# MODEL_NAME = "claude-3-5-sonnet-20240620"

# Function to generate screenplay using Anthropic Claude 3.5 (Messages API)
def generate_screenplay(data):
    user_prompt = f"""Generate a short 3-4 minute screenplay for a children's short firlm based on the following details. The output screenplay must be formatted as a JSON.
    <input>
    Here are the details
    Theme: {data['theme']}
    Story Summary: {data['story_summary']}
    Characters: {data['characters']}
    Props: {data['props']}
    VFX: {'Yes' if data['vfx'] else 'No'}
    Drones: {'Yes' if data['drones'] else 'No'}
    Location: {data['location']}
    </input>

    <guidelines>
    The JSON output must be in the following JSON format. You can use up to 5 scenes (for scene_number), and up to 3 shots (for shot_number) per scene. Make the spoke_lines short and concise.
    {{
        "title": "Title of the screenplay",
        "synopsis": "3 to 4 sentence synopsis of the screenplay",
        "characters": "Characters in the screenplay, their names (if provided), and roles description",
        "scenes": [
            {{
                "scene_number": "Scene number (1 to 5)",
                "scene_breakdown": "Breakdown of the scene, including location, characters, props, and what's happening",
                "scene_props": "Props needed for this scene",
                "shots": [
                    {{
                        "shot_number": "Shot number (1 to 3)",
                        "shot_description": "Description of the shot, including camera movement and character tips",
                        "spoken_lines": "Spoken lines for the shot"
                    }}
                ]
            }}
        ],
        "included_props": "suggested props and wardrobe used in the screenplay" ,
        "suggested_props": "suggested props and wardrobe that can be used in the screenplay but we are not using it in the screenplay"
    }}

    An example of a title is: 
        Knights of the Round Table
    An example of a synopsis is: 
        A whimsical adventure where cardboard swords meet courageous hearts! Join our tiny knights as they face dragons, solve puzzles, and navigate magical corridors to reclaim their stolen crown. Prepare for a journey full of laughter and adventure, where courage comes in all sizes!
    An example of characters is: 
        John - Hero: A young man with traditional dressing and a swad \nAnna - Friend: A girl with blod hair dressed in pink dress and is very sweet and pretty \nDragon - Villain: Huge red and orange color dragon with fire powers.
    An example of a scene_breakdown is: 
        Scene 1: INT. FAMILY LIVING ROOM - DAY \n The living room has been transformed into a makeshift castle, with towers made of cardboard and a round table crafted from cardboard on an ottoman. Characters stand around the round table in homemade knight's armor, each holding a cardboard sword
    An example of a scene_props is:
        Sword, Dress, Round Table
    An example of a shot_description is: 
        Wide shot to establish the imaginative setting
    An example of spoken lines is: 
        Character 1: (placing his sword over the table with dramatic flair) "Knights of the Round Table, our quest begins! We face great challenges today!"
        Character 2: (pumping his fist) "We shall retrieve our lost crown!"
        Character 3: (jumping slightly, excited) "And rescue the mythical creature from dire perils!"
        Character 4: (raising sword) " No dragon shall breach our castle walls!"

    Double check the JSON format to make sure the keys and formatting are correct. Use double quotes for all string values. Use single quotes for spoken_lines.This is very important.
    </guidelines>
    """

    message_list = [
        {
            "role": 'user',
            "content": user_prompt
        },
        {
            "role": "assistant",
            "content": "Here is the JSON requested:\n{"
        }
    ]

    response = anthropic_client.messages.create(
        model="claude-3-5-sonnet-20240620",
        max_tokens=8192,
        messages=message_list
    )

    output = response.content[0].text
    print(output)

    return output

# Function to generate an image using OpenAI DALL-E
def generate_image(prompt):
    response = open_ai_client.images.generate(
        model = "dall-e-3",
        prompt=prompt,
        n=1,
        size="1024x1024"
    )

    # Get the image URL correctly using dot notation
    image_url = response.data[0].url  # Access the attribute

    # Download the image and save it as a .jpeg
    image_data = requests.get(image_url).content
    image_io = BytesIO(image_data)  # Create a BytesIO object

    return image_io

def json_to_string(json_obj):
    result = []
    
    def traverse(obj, parent_key=None):
        if isinstance(obj, dict):
            for key, value in obj.items():
                traverse(value, key)
        elif isinstance(obj, list):
            for item in obj:
                traverse(item, parent_key)
        else:
            if parent_key == "scene_number":
                result.append(f"Scene Number: {obj}")
            elif parent_key == "shot_number":
                result.append(f"Shot Number: {obj}")
            else:
                result.append(str(obj))
    
    traverse(json_obj)
    return '\n\n'.join(result)

def create_docx(screenplay_json, image):
    doc = Document()

    # Add logo to footer
    section = doc.sections[0]
    section.footer_distance = Inches(0.1)
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer_run = footer_para.add_run()
    #footer_run.add_picture('logo.png', width=Inches(1))  # Adjust path and size as needed

    # Adjust spacing to move logo to bottom
    #footer_para.paragraph_format.space_before = Pt(0)
    #footer_para.paragraph_format.space_after = Pt(0)

    # Add title
    title = screenplay_json["title"]
    doc.add_heading(title, 0)
    # Add the image to the document
    doc.add_picture(image, width=Inches(6.0))

    doc.add_page_break()

    #Second page: Welcome and tips
    doc.add_heading("Welcome to your Personalized KidVentures Video", level=1)
    paragraph1 = doc.add_paragraph()
    run1 = paragraph1.add_run("Thank you for choosing KidVenture Studios! ")
    run1.bold = True
    paragraph1.add_run("🎬✨ We’re thrilled to bring your vision to life with a script designed especially for YOU. Follow the steps below to ensure a stress-free and fun-filled experience, packed with action, drama, or maybe a few dino-roars 🦖 (we won’t judge). So grab your camera, call in your co-stars (even if it's just the family dog 🐾), and get ready to shout, “Action!”")
    #doc.add_paragraph("**Thank you for choosing KidVenture Studios!** 🎬✨ We’re thrilled to bring your vision to life with a script designed especially for YOU.  Follow the steps below to ensure a stress-free and fun-filled experience, packed with action, drama, or maybe a few dino-roars 🦖 (we won’t judge). So grab your camera, call in your co-stars (even if it's just the family dog 🐾), and get ready to shout, “Action!")
    doc.add_heading("How to Use", level=1)
    doc.add_paragraph("Use this script as a guide from top-to-bottom for your children's adventure film. Encourage improvisation and let the kids add their own ideas. Focus on capturing the fun and imagination rather than perfect line delivery.")
    doc.add_heading("Helpful Tips", level=1)
    # Define the content with bold markers ** for emphasis - for Tips bullet
    content = (
        "- **Use a SMARTPHONE** in LANDSCAPE mode for filming\n"
        "- **Use natural light** - film outdoors or near windows whenever possible\n"
        "- **Steady your shots** - keep your camera still with steady movement and avoid shake\n"
        "- **Clean your lens** - remove any unintended debris or dirt\n"
        "- **Mind your sound** - minimize background noise and speak lines clearly\n"
        "- **Frame subject** - ensure the person or item of importance is within shot, not cut off\n"
        "- **Save each clip** with scene and shot number, ie- SCENE1_SHOT1"
    )
    # Create a new paragraph
    paragraph2 = doc.add_paragraph()
    # Split the content by ** and format accordingly
    parts = content.split("**")
    for i, part in enumerate(parts):
        if i % 2 == 1:  # Odd indexes are the parts to be bolded
            run2 = paragraph2.add_run(part)
            run2.bold = True
        else:  # Even indexes are the regular text
            paragraph2.add_run(part)
    doc.add_page_break()

    # Third page: Synopsis and Characters
    doc.add_heading("Your Story", level=1)
    doc.add_heading("Synopsis", level=1)
    doc.add_paragraph(screenplay_json["synopsis"])
    doc.add_heading("Characters", level=1)
    doc.add_paragraph(screenplay_json["characters"])
    doc.add_heading("Props", level=1)
    #doc.add_paragraph(screenplay_json["suggested_props"])
    # Extracting included and suggested props from the screenplay_json
    included_props_list = screenplay_json["included_props"].split(',')
    suggested_props_list = screenplay_json["suggested_props"].split(',')
    # Create a table with 0 initial rows (since we'll be dynamically adding rows)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    # Add "Included" props if there are any
    if included_props_list:
        for prop in included_props_list:
            row_cells = table.add_row().cells
            row_cells[1].text = prop.strip()  # Add prop name in the second column
        # Merge the first column cells for "Included" props
        included_cell = table.cell(0, 0).merge(table.cell(len(included_props_list) - 1, 0))
        included_cell.text = "Included"  # Set the merged cell text for Included

    # Add "Suggested" props if there are any
    if suggested_props_list:
        start_row = len(included_props_list)  # Starting point for Suggested props
        for prop in suggested_props_list:
            row_cells = table.add_row().cells
            row_cells[1].text = prop.strip()  # Add prop name in the second column
        # Merge the first column cells for "Suggested" props
        suggested_cell = table.cell(start_row, 0).merge(table.cell(len(table.rows) - 1, 0))
        suggested_cell.text = "Suggested"  # Set the merged cell text for Suggested


    doc.add_page_break()

    # Remaining pages: Scenes and additional information
    doc.add_heading("Scenes", level=1)
    for scene in screenplay_json["scenes"]:
        scene_heading = f"Scene {scene['scene_number']}"
        doc.add_heading(scene_heading, level=2)
        doc.add_paragraph(scene["scene_breakdown"])
        paragraph3 = doc.add_paragraph()
        run3 = paragraph3.add_run("Props needed:\n")
        run3.bold = True
        
        # Check if 'scene_props' exists and is not None
        scene_props = scene.get("scene_props", "No specific props listed.")
        paragraph3.add_run(scene_props)

        for shot in scene["shots"]:
            shot_heading = f"Shot {shot['shot_number']}"
            # Add heading and then indent it
            heading = doc.add_heading(shot_heading, level=3)
            heading.paragraph_format.left_indent = Inches(0.5)  # Indent by 0.5 inches
            # Add shot description and indent it
            shot_description = doc.add_paragraph(shot["shot_description"])
            shot_description.paragraph_format.left_indent = Inches(0.5)  # Indent by 0.5 inches
            # Add spoken lines and indent them
            spoken_lines = doc.add_paragraph(shot["spoken_lines"], style='Quote')
            spoken_lines.paragraph_format.left_indent = Inches(0.5)  # Indent by 0.5 inches

    # # Additional sections
    # doc.add_heading("Suggested Props", level=1)
    # doc.add_paragraph(screenplay_json["suggested_props"])

    # doc.add_heading("How to Use", level=1)
    # doc.add_paragraph(screenplay_json["how_to_use"])

    # doc.add_heading("Filming Tips", level=1)
    # doc.add_paragraph(screenplay_json["filming_tips"])

    # doc.add_heading("Trendy Shots", level=1)
    # doc.add_paragraph(screenplay_json["trendy_shots"])

    # Save the document to a BytesIO object to make it downloadable
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)  # Reset the pointer to the beginning of the file

    return doc_io

# Streamlit Interface

st.title("🎬 Kidventure Studios")
st.subheader("Screenplay Generator")

# Initialize trackers
if "title" not in st.session_state:
    st.session_state.title = None
if "screenplay" not in st.session_state:
    st.session_state.screenplay = None
if "screenplay_text" not in st.session_state:
    st.session_state.screenplay_text = None    
if "image" not in st.session_state:
    st.session_state.image = None

screenplay = None
generated_image = None

col1, col2 = st.columns([1,1])

with col1:
    st.header("Screenplay Inputs")
    theme_options = ["Adventure", "Fantasy", "Mystery", "Action", "Sports", "Other"]
    selected_theme = st.selectbox("Theme", options=theme_options)
    if selected_theme == "Other":
        custom_theme = st.text_input("Custom Theme")
        theme = custom_theme if custom_theme else "Custom"
    else:
        theme = selected_theme

    story_summary = st.text_area("Story Summary", placeholder="A brief summary of the story")
    characters = st.text_area("Characters", placeholder="List of characters (e.g. John - Hero, Anna - Friend)")
    props = st.text_area("Props and Wardrobe", placeholder="Describe props and wardrobe")

    vfx = st.checkbox("Include VFX shots")
    drones = st.checkbox("Include Drone shots")

    selected_location = st.selectbox("Location", options=["Home", "Other"])
    if selected_location == "Other":
        custom_location = st.text_input("Custom Location (if other)")
        location = custom_location
    else:
        location = selected_location

    generate_image_checkbox = st.checkbox("Generate Cover Image")

    if st.button('Generate Screenplay'):
        with st.spinner('Generating screenplay... Please wait.'):
            data = {
                "theme": theme,
                "story_summary": story_summary,
                "characters": characters,
                "props": props,
                "vfx": vfx,
                "drones": drones,
                "location": custom_location if location == "Other" else location
            }
            screenplay = generate_screenplay(data)

            screenplay_json = json.loads("{" + screenplay[:screenplay.rfind("}") + 1])
            
            screenplay_text = json_to_string(screenplay_json)
            
            title = screenplay_json["title"]

            st.session_state.title = title
            st.session_state.screenplay = screenplay_json
            st.session_state.screenplay_text = screenplay_text

            # Generate the image based on the theme
            # IF checkbox do this
            if generate_image_checkbox:  # Use the new variable name here
                # title, characters, props, synopsis = extract_image_details(screenplay)
                image_prompt = (
                    f"Make a {theme} movie poster for the following screenplay, in the style of {theme} theme film (not a cartoon)."
                    f"Title: {title}"
                    f"Synopsis: {story_summary}"
                    f"Characters: {characters}"
                    f"highlighting the theme of the screenplay. NEVER PUT TEXT ON THE MOVIE POSTER."
                )
                st.session_state.image_prompt = image_prompt
                image_path = generate_image(image_prompt)
                #image_io = generate_image(image_prompt)  # Now returns BytesIO object
                st.session_state.image = image_path
                # st.image(image_path, caption="Generated Cover Image", use_column_width=True)

            # st.subheader("Generated Screenplay")
            # st.text_area("Screenplay", value=screenplay, height=400)

            # # Create a .docx file from the screenplay text
            # docx_file = create_docx(screenplay,image_path)

            # # Option to download screenplay as .docx
            # st.download_button(
            #     label="Download Screenplay",
            #     data=docx_file,
            #     file_name="screenplay.docx",
            #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            # )
with col2:
    st.header("Generated Screenplay")
    screenplay_container = st.container(height=380)
    with screenplay_container:
        if st.session_state.screenplay_text:
            st.write("### Generated Document")
            st.text_area("Screenplay", value=st.session_state.screenplay_text, height=380)

    if 'images' not in st.session_state:
        st.session_state.images = []

    if 'regeneration_count' not in st.session_state:
        st.session_state.regeneration_count = 0

    cover_image_container = st.container(height=380)
    with cover_image_container:
        # st.header("Generated Cover Image")
        if st.session_state.image:
            st.write("### Generated Cover Image")
            st.image(st.session_state.image, caption="Generated Cover Image", use_column_width=True)

    if st.session_state.image:
        if st.button("Regenerate Cover Image"):
            with st.spinner('Regenerating Cover Image... Please wait.'):
                # Generate a new image based on the prompt
                image_path = generate_image(st.session_state.image_prompt)
                # Append the new image to the session state list
                st.session_state.images.append(image_path)
                # Increment the regeneration count
                st.session_state.regeneration_count += 1
                # Display the newly generated image immediately with a unique caption
                with cover_image_container:
                    for path in st.session_state.images:
                        st.image(path, caption=f"Regenerated Cover Image {st.session_state.regeneration_count}", use_column_width=True)
    
    if st.session_state.screenplay and st.session_state.image:  
        # Create a .docx file from the screenplay text
        docx_file = create_docx(st.session_state.screenplay,st.session_state.image)

        # Text input for the user to enter the file name
        file_name = st.text_input("Enter the name of your screenplay", placeholder="Enter the name of the file (e.g., my_play)")
        # Ensure the file name ends with .docx
        if not file_name.endswith(".docx"):
            file_name += ".docx"

        # Option to download screenplay as .docx
        st.download_button(
            label="Download Screenplay",
            data=docx_file,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )